import statistics
from .models import Test, Plate,Well
from docx import Document
from pathlib import Path
import os
from django.http import Http404, HttpResponse

class Formula():

    def __init__(self,t:Test):
        self.test = t

        self.wellsDict = {}
        self.statsDict = {}
        
        for p in Plate.objects.filter(test = self.test):
            wells = Well.objects.filter(plate = p)
            dos = wells[0].dosage
            counter = 1
            sec = []
         
            self.statsDict[p.num] = {"sec"+str(counter): {}}
            stats = self.statsDict[p.num]
            self.wellsDict[p.num] = {"sec"+str(counter): []}
            dic = self.wellsDict[p.num]
            first_avg = None

            for w in wells:
                if(w.dosage == dos):
                    sec.append(w)
                    dic["sec"+str(counter)] =  []
                    dic["sec"+str(counter)] = sec
                    dos = w.dosage
                   
                    if(counter == 4):
                        
                        stats["sec"+str(counter)]["avg"] = self.average(sec)
                        stats["sec"+str(counter)]["st_dev"] = self.st_dev(sec)
                        stats["sec"+str(counter)]["si"]  =  stats["sec"+str(counter)]["avg"]/first_avg
                       


                else:
                    stats["sec"+str(counter)]["avg"] = self.average(sec)
                    stats["sec"+str(counter)]["st_dev"] = self.st_dev(sec)
                    if(counter >1):
                        stats["sec"+str(counter)]["si"]  =  stats["sec"+str(counter)]["avg"]/first_avg
                    else:
                        first_avg = stats["sec"+str(counter)]["avg"]

                    counter = counter +1
                    sec = []
                    sec.append(w)
                    stats["sec"+str(counter)] = {}
                    dic["sec"+str(counter)] =  []
                    dic["sec"+str(counter)] = sec
                    dos = w.dosage
    
        self.resultDict = {
            "4D" :
        {"BeSO4 100uM [1X10-4]": self.statsDict[1]["sec2"]["si"], "BeSO4 10uM [1X10-5]": self.statsDict[1]["sec3"]["si"] ,"BeSO4 1uM [1X10-6]": self.statsDict[1]["sec4"]["si"]},
       
        "6D" :
        {"BeSO4 100uM [1X10-4]": self.statsDict[3]["sec2"]["si"], "BeSO4 10uM [1X10-5]": self.statsDict[3]["sec3"]["si"],"BeSO4 1uM [1X10-6]": self.statsDict[3]["sec4"]["si"]}
        }

        self.controlDict ={
            "4-PHA" : self.controlStat(self.statsDict[2]) ,
            "6-PWM" : self.controlStat(self.statsDict[4])
        }
    
    def controlStat(self, s :dict):
        ls = []
        counter = 1
        for key,value in s.items():
            print(key,value)
            if counter>1:
             ls.append(value["si"])
            counter = counter +1
        return max(ls)



    def average(self,lst):
        nums = []
        for w in lst:
            nums.append(w.reading)

        return sum(nums) / len(nums)

    def st_dev(self,lst):
        nums = []
        for w in lst:
            nums.append(w.reading)

        return statistics.stdev(nums)

    def exportToDoc(self,isNormal,isPositive,profession):
        path = Path(__file__).resolve().parent
        docStr = str(path)+"/base_xls/system_files/template.docx"
        newStr = str(path)+"/exports/" + str(self.test.patient) +" " + str(self.test.patient.id) + "RES.docx"
        doc = Document(docStr)
        dic = {"name":self.test.patient.__str__(), "date":self.test.date.__str__(),"id":self.test.patient.id,"profession":profession,
        "profession" : profession,"phaResult":self.controlDict["4-PHA"],"pwmResult":self.controlDict["6-PWM"],
        "4-4Result":self.resultDict["4D"]["BeSO4 100uM [1X10-4]"],"4-5Result":self.resultDict["4D"]["BeSO4 10uM [1X10-5]"],"4-6Result":self.resultDict["4D"]["BeSO4 1uM [1X10-6]"],
        "6-4Result":self.resultDict["6D"]["BeSO4 100uM [1X10-4]"],"6-5Result":self.resultDict["6D"]["BeSO4 10uM [1X10-5]"],"6-6Result":self.resultDict["6D"]["BeSO4 1uM [1X10-6]"]}
        if(isPositive):
            dic["isPositiveResult"] = "חיובי"
        else:
            dic["isPositiveResult"] = "שלילי"
        if(isNormal):
            dic["isNormalResult"] = "כן"
        else:
            dic["isNormalResult"] = "לא"  

        for key in dic:
            findReplace(doc,key,str(dic[key]))
            doc.save(newStr)
        
        file_path = newStr
        if os.path.exists(file_path):
            with open(file_path, 'rb') as fh:
                response = HttpResponse(fh.read(), content_type="application/vnd.ms-word")
                response['Content-Disposition'] = 'inline; filename=' + os.path.basename(file_path)
                return response
        raise Http404
        

def findReplace(document,str,rplc):
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if str in paragraph.text:
                        paragraph.text = paragraph.text.replace(str, rplc)


